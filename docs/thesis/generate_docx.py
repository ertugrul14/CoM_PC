"""Convert curbside_intensification_thesis.md to a styled DOCX."""
import pathlib
import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

SRC = pathlib.Path(__file__).parent / "curbside_intensification_thesis.md"
DST = pathlib.Path(__file__).parent / "curbside_intensification_thesis.docx"


def setup_styles(doc: Document):
    """Configure document styles for academic formatting."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Georgia"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    p.paragraph_format.space_before = Pt(120)
    p.paragraph_format.space_after = Pt(12)


def add_subtitle(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    p.paragraph_format.space_after = Pt(24)


def add_heading(doc: Document, text: str, level: int):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    if level == 1:
        h.paragraph_format.page_break_before = True


def add_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment

    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|\^[0-9]+)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.font.italic = True
        elif re.match(r'\^\d+', part):
            run = p.add_run(part[1:])
            run.font.superscript = True
            run.font.size = Pt(8)
        else:
            run = p.add_run(part)
            if bold:
                run.font.bold = True
            if italic:
                run.font.italic = True
    return p


def add_figure_placeholder(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_table(doc: Document, header_row: list, data_rows: list):
    table = doc.add_table(rows=1 + len(data_rows), cols=len(header_row))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, cell_text in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = cell_text.strip()
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    for row_idx, row_data in enumerate(data_rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_text.strip()
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)


def parse_table_line(line: str) -> list:
    """Parse a markdown table row into cells."""
    cells = [c.strip() for c in line.split("|")]
    return [c for c in cells if c and not re.match(r'^[-:]+$', c)]


def is_separator_row(line: str) -> bool:
    """Check if a line is a markdown table separator."""
    return bool(re.match(r'^\s*\|[\s\-:|]+\|\s*$', line))


def main():
    doc = Document()
    setup_styles(doc)

    md_text = SRC.read_text(encoding="utf-8")
    lines = md_text.split("\n")

    i = 0
    first_h2 = True
    in_footnotes = False

    while i < len(lines):
        line = lines[i]

        # Skip horizontal rules
        if re.match(r'^---+\s*$', line):
            i += 1
            continue

        # H1 title
        if line.startswith("# ") and not line.startswith("## "):
            add_title(doc, line[2:].strip())
            i += 1
            continue

        # H2
        if line.startswith("## "):
            text = line[3:].strip()
            if first_h2:
                add_subtitle(doc, text)
                first_h2 = False
            else:
                if text == "Footnotes":
                    in_footnotes = True
                    add_heading(doc, text, level=1)
                else:
                    add_heading(doc, text, level=1)
            i += 1
            continue

        # H3
        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), level=2)
            i += 1
            continue

        # Figure placeholder
        if line.startswith("[Figure"):
            add_figure_placeholder(doc, line)
            i += 1
            continue

        # Table detection
        if "|" in line and i + 1 < len(lines) and is_separator_row(lines[i + 1]):
            header = parse_table_line(line)
            i += 2  # skip header and separator
            data_rows = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                data_rows.append(parse_table_line(lines[i]))
                i += 1
            if header and data_rows:
                add_table(doc, header, data_rows)
            continue

        # Bullet points
        if line.startswith("- ") or line.startswith("* "):
            text = line[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|\^[0-9]+)', text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                elif part.startswith("*") and part.endswith("*"):
                    run = p.add_run(part[1:-1])
                    run.font.italic = True
                elif re.match(r'\^\d+', part):
                    run = p.add_run(part[1:])
                    run.font.superscript = True
                    run.font.size = Pt(8)
                else:
                    p.add_run(part)
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line).strip()
            p = doc.add_paragraph(style="List Number")
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|\^[0-9]+)', text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                elif part.startswith("*") and part.endswith("*"):
                    run = p.add_run(part[1:-1])
                    run.font.italic = True
                elif re.match(r'\^\d+', part):
                    run = p.add_run(part[1:])
                    run.font.superscript = True
                    run.font.size = Pt(8)
                else:
                    p.add_run(part)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph (may span multiple lines until blank)
        para_lines = []
        while i < len(lines) and lines[i].strip() and not lines[i].startswith("#") \
                and not lines[i].startswith("[Figure") and not lines[i].startswith("---") \
                and not lines[i].startswith("- ") and not lines[i].startswith("* ") \
                and not re.match(r'^\d+\.\s', lines[i]) \
                and not ("|" in lines[i] and i + 1 < len(lines) and is_separator_row(lines[i + 1])):
            para_lines.append(lines[i])
            i += 1

        if para_lines:
            full_text = " ".join(para_lines)
            add_paragraph(doc, full_text)

    doc.save(str(DST))
    size_kb = DST.stat().st_size / 1024
    print(f"DOCX generated: {DST} ({size_kb:.0f} KB)")


if __name__ == "__main__":
    main()
